import streamlit as st
import json
import re
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from Login import require_login

DATA_DIR = Path("data")
OUTPUT_DIR = Path("output")
WEEK = (datetime.today() - timedelta(days=7)).isocalendar()[1]
vandaag = datetime.today()
MA = vandaag - timedelta(days=vandaag.weekday())
MA1 = MA - timedelta(weeks=1)
ZO1 = MA1 + timedelta(days=6)

st.set_page_config(page_title="THOR Stedelijk Informatiebeeld", layout="wide")

require_login()

stadsdelen = ["Algemeen beeld", "Centrum", "Noord", "Oost", "Zuid", "Zuidoost", "Weesp", "West", "Nieuw-West", "VOV", "Nautisch toezicht"]
onderdelen = ["Overlast personen", "Overlast jeugd", "Afval", "Parkeeroverlast/verkeersoverlast", "Overige reguliere taken"]
nautisch = ["Incidenten", "Regulier Werk", "CityControl", "SIG-meldingen"]

st.title(f"Invoer Stedelijk Beeld Week {WEEK}")

stadsdeel = st.selectbox("Stadsdeel / Specialisme / Algemeen Beeld", stadsdelen)

with st.form("invoer_form"):
    teksten = {}

    # Laad eerder gesubmitte data voor het gekozen stadsdeel
    for onderdeel in nautisch if stadsdeel == "Nautisch Toezicht" else onderdelen:
        safe_onderdeel = re.sub(r"[\\/]", "_", onderdeel)
        filename = f"{WEEK}_{safe_onderdeel}_{stadsdeel}.json".replace(" ", "_")
        tekst = ""
        filepath = DATA_DIR / filename
        if filepath.exists():
            with open(filepath, encoding="utf-8") as f:
                data_json = json.load(f)
                tekst = data_json.get("tekst", "")
        teksten[onderdeel] = st.text_area(f"{onderdeel}", value=tekst, height=100)

    submitted = st.form_submit_button("Opslaan")

    if submitted:
        DATA_DIR.mkdir(exist_ok=True)
        for onderdeel, tekst in teksten.items():
            safe_onderdeel = re.sub(r"[\\/]", "_", onderdeel)  # slash vervangen voor bestandsnaam
            filename = f"{WEEK}_{safe_onderdeel}_{stadsdeel}.json".replace(" ", "_")
            with open(DATA_DIR / filename, "w", encoding="utf-8") as f:
                json.dump({
                    "week": WEEK,
                    "onderdeel": onderdeel,
                    "stadsdeel": stadsdeel,
                    "tekst": tekst
                }, f)
        st.success(f"Invoer opgeslagen voor {stadsdeel}")

if st.button("Genereer Word rapport"):
    if not DATA_DIR.exists() or not any(DATA_DIR.glob(f"{WEEK}_*.json")):
        st.warning("Geen invoer gevonden voor deze week.")
    else:
        doc = Document()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Amsterdam Sans'

        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        doc.add_paragraph()
        titel = doc.add_paragraph()
        titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = titel.add_run(f"Informatiebeeld THOR \n Week {WEEK} \n {MA1.strftime('%d-%m-%Y')} t/m {ZO1.strftime('%d-%m-%Y')}")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Amsterdam Sans'
        titel.paragraph_format.space_after = Pt(60)

        datum = doc.add_paragraph()
        datum.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date = datetime.now().strftime('%d-%m-%Y')
        datum_run = datum.add_run(f"Team Informatiemanagement THOR \n {date}")
        datum_run.font.size = Pt(26)
        datum_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rood
        datum_run.italic = True
        datum_run.bold = True
        datum_run.font.name = 'Amsterdam Sans'
        datum.paragraph_format.space_after = Pt(60)

        doc.add_page_break()

        # Inhoudsopgave
        inhoud = doc.add_heading(level=1)
        inhoud_run = inhoud.add_run("Inhoud")
        inhoud.paragraph_format.space_after = Pt(12)
        inhoud_run.font.color.rgb = RGBColor(0, 0, 0)
        inhoud_run.font.name = 'Amsterdam Sans'

        weekbeeld = doc.add_paragraph()
        weekbeeld_run = weekbeeld.add_run("1. Weekbeeld per thema")
        weekbeeld_run.bold = True
        weekbeeld.style = 'Normal'
        weekbeeld.paragraph_format.left_indent = Pt(16)
        
        # Subitems
        inhoudsopgave_lijst = onderdelen + ["Nautisch Toezicht"]
        for j, item in enumerate(inhoudsopgave_lijst, start=1):
            subnummer = f"1.{j}"
            para = doc.add_paragraph(f"{subnummer} {item}")
            para.style = 'Normal'
            para.paragraph_format.left_indent = Pt(36)
        
        weekbeeld = doc.add_paragraph()
        weekbeeld_run = weekbeeld.add_run("2. Weekbeeld in cijfers")
        weekbeeld_run.bold = True
        weekbeeld.style = 'Normal'
        weekbeeld.paragraph_format.left_indent = Pt(16)

        doc.add_page_break()

        # Nieuwe sectie maken met 2 kolommen
        new_section = doc.add_section()
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height
        new_section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), '2')

        heading = doc.add_heading(level=1)
        weekbeeld_run = heading.add_run("1. Weekbeeld per thema")
        weekbeeld_run.bold = True
        weekbeeld_run.font.size = Pt(20)
        weekbeeld_run.font.name = 'Amsterdam Sans'
        weekbeeld_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        # Data laden
        data = {}
        for file in DATA_DIR.glob(f"{WEEK}_*.json"):
            with open(file, encoding="utf-8") as f:
                entry = json.load(f)
                data.setdefault(entry["onderdeel"], {})[entry["stadsdeel"]] = entry["tekst"]

        # Nummer teller starten
        counter = 1.1

        # Hoofdonderdelen met oplopend nummer
        for idx, onderdeel in enumerate(onderdelen, start=1):
            nummer = f"1.{idx}"
            heading = doc.add_heading(level=2)
            run = heading.add_run(f"{nummer} {onderdeel}")
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rood
            run.font.name = 'Amsterdam Sans'
            run.font.size = Pt(10)

            counter += 0.1

            stadsdeel_data = data.get(onderdeel, {})
            for stadsdeel in stadsdelen:
                if stadsdeel != "Nautisch toezicht":
                    heading2 = doc.add_heading(level=3)
                    run2 = heading2.add_run(stadsdeel)
                    run2.bold = True
                    run2.font.name = 'Amsterdam Sans'
                    run2.font.size = Pt(10)

                    if stadsdeel == "Algemeen beeld":
                        run2.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)  # blauw
                    else:
                        run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # zwart

                    tekst = stadsdeel_data.get(stadsdeel, "")
                    if tekst:
                        for para in tekst.split('\n'):
                            p = doc.add_paragraph(para)
                            p.style = 'Normal'
                            for run in p.runs:
                                run.font.name = 'Amsterdam Sans'
                                run.font.size = Pt(10)

        # Nautisch Toezicht als laatste met nummer
        heading = doc.add_heading(level=2)
        run = heading.add_run(f"1.{len(onderdelen)+1} Nautisch toezicht")
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rood
        run.font.name = 'Amsterdam Sans'
        run.font.size = Pt(10)

        counter += 0.1

        for nautisch_onderdeel in nautisch:
            heading2 = doc.add_heading(level=3)
            run2 = heading2.add_run(nautisch_onderdeel)
            run2.bold = True
            run2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # zwart
            run2.font.name = 'Amsterdam Sans'
            run2.font.size = Pt(10)

            nautisch_data = data.get(nautisch_onderdeel, {})
            tekst = nautisch_data.get("Nautisch toezicht", "")
            if tekst:
                for para in tekst.split('\n'):
                    p = doc.add_paragraph(para)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Amsterdam Sans'
                        run.font.size = Pt(10)
        
        heading = doc.add_heading(level=1)
        weekbeeld_run = heading.add_run("2. Weekbeeld in cijfers")
        weekbeeld_run.bold = True
        weekbeeld_run.font.size = Pt(20)
        weekbeeld_run.font.name = 'Amsterdam Sans'
        weekbeeld_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        OUTPUT_DIR.mkdir(exist_ok=True)
        output_path = OUTPUT_DIR / f"{vandaag} THOR Stedelijk Informatiebeeld week {WEEK}.docx"
        doc.save(output_path)

        st.success(f"Word rapport gegenereerd: {output_path}")
        with open(output_path, "rb") as f:
            st.download_button("Download Word rapport", f, file_name=output_path.name)
